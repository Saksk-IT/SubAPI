#!/usr/bin/env python3
"""生成可导入飞书、且图片真正内嵌的 Word 教程。"""

from __future__ import annotations

import argparse
import base64
import io
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.image.image import Image
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
except ImportError as error:  # pragma: no cover - 命令行环境缺依赖时使用
    raise SystemExit(
        "缺少 python-docx；请使用 Codex 文档运行时或安装 python-docx>=1.2。"
    ) from error

if __package__:
    from . import export_feishu_guides as markdown_exporter
    from . import export_guides_v2
    from . import export_word_guides_v2 as word_v2
    from .export_word_guides_io import (
        check_documents,
        export_documents,
        validated_docx_path,
    )
else:
    import export_feishu_guides as markdown_exporter
    import export_guides_v2
    import export_word_guides_v2 as word_v2
    from export_word_guides_io import (
        check_documents,
        export_documents,
        validated_docx_path,
    )


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT_DIR = REPO_ROOT / "docs" / "static-guides" / "feishu-word"
FIXED_PACKAGE_TIME = (2026, 7, 10, 0, 0, 0)
FIXED_CORE_TIME = datetime(2026, 7, 10, tzinfo=timezone.utc)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
BODY_FONT = "Arial"
EAST_ASIA_FONT = "PingFang SC"
MONO_FONT = "Menlo"
ACCENT_BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK_BLUE = RGBColor(11, 37, 69)
MUTED_GRAY = RGBColor(95, 99, 104)
LIGHT_BLUE_FILL = "E8EEF5"
LIGHT_GRAY_FILL = "F4F6F9"
LINK_BLUE = "0563C1"
IMAGE_PATTERN = re.compile(r"^!\[([^\]]*)\]\(([^)\n]+)\)$")
HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")
ORDERED_PATTERN = re.compile(r"^\s*\d+\.\s+(.+)$")
BULLET_PATTERN = re.compile(r"^\s*-\s+(.+)$")
TABLE_SEPARATOR_PATTERN = re.compile(
    r"^\|\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|$"
)
INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|<https?://[^>]+>)"
)
RFONTS_ELEMENT_PATTERN = re.compile(rb"<w:rFonts\b[^>]*>")
RFONTS_ATTRIBUTE_PATTERN = re.compile(
    rb"\bw:(ascii|hAnsi|eastAsia|cs)=(['\"])(.*?)\2"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="输出目录（默认随 edition 选择 V1 或 V2 成品目录）",
    )
    parser.add_argument(
        "--edition",
        choices=("v1", "v2"),
        default="v1",
        help="导出版本（默认：v1）",
    )
    parser.add_argument(
        "--check",
        action="store_true",
        help="只校验现有 Word 产物是否与当前源稿一致",
    )
    return parser.parse_args()


def set_run_font(
    run,
    *,
    name: str = BODY_FONT,
    east_asia: str = EAST_ASIA_FONT,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style(style, *, size: float, color: RGBColor, before: float, after: float) -> None:
    style.font.name = BODY_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = True
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:cs"), EAST_ASIA_FONT)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True


def configure_document(doc: Document, title: str, is_parent: bool) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:cs"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    configure_style(
        doc.styles["Heading 1"],
        size=16,
        color=ACCENT_BLUE,
        before=18,
        after=10,
    )
    configure_style(
        doc.styles["Heading 2"],
        size=13,
        color=ACCENT_BLUE,
        before=14,
        after=7,
    )
    configure_style(
        doc.styles["Heading 3"],
        size=12,
        color=DARK_BLUE,
        before=10,
        after=5,
    )

    caption = doc.styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = MUTED_GRAY
    caption._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    caption._element.get_or_add_rPr().rFonts.set(qn("w:cs"), EAST_ASIA_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.keep_together = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run(
        f"SAK AI · {'父教程' if is_parent else '客户端子教程'} · {title}"
    )
    set_run_font(header_run, size=8.5, color=MUTED_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("第 ")
    set_run_font(footer_run, size=8.5, color=MUTED_GRAY)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    footer_run._r.extend((field_begin, instruction, field_end))
    suffix_run = footer.add_run(" 页")
    set_run_font(suffix_run, size=8.5, color=MUTED_GRAY)

    properties = doc.core_properties
    properties.title = title
    properties.subject = "飞书导入版配置教程"
    properties.author = "SubAPI"
    properties.last_modified_by = "SubAPI"
    properties.created = FIXED_CORE_TIME
    properties.modified = FIXED_CORE_TIME


def add_cover(doc: Document, title: str, is_parent: bool) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(64)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    kicker_run = kicker.add_run("飞书导入版教程")
    set_run_font(kicker_run, size=10.5, color=ACCENT_BLUE, bold=True)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(10)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title)
    set_run_font(title_run, size=27, color=INK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(54)
    subtitle_run = subtitle.add_run(
        "中转注册与密钥配置父教程"
        if is_parent
        else "中转注册与密钥配置的客户端子教程"
    )
    set_run_font(subtitle_run, size=13, color=MUTED_GRAY)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("所有教程截图均已嵌入 Word 文件内部")
    set_run_font(note_run, size=10, color=ACCENT_BLUE, bold=True)
    doc.add_page_break()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    fonts.set(qn("w:cs"), EAST_ASIA_FONT)
    run_properties.extend((fonts, color, underline))
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend((run_properties, text_element))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_text(paragraph, text: str, *, bold: bool = False) -> None:
    line_break_parts = re.split(r"<br\s*/?>", text, flags=re.IGNORECASE)
    if len(line_break_parts) > 1:
        for part_index, part in enumerate(line_break_parts):
            add_inline_text(paragraph, part, bold=bold)
            if part_index < len(line_break_parts) - 1:
                paragraph.add_run().add_break()
        return
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            plain_run = paragraph.add_run(
                text[position : match.start()].replace(r"\*", "*")
            )
            set_run_font(plain_run, bold=bold)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                name=MONO_FONT,
                east_asia=EAST_ASIA_FONT,
                size=10,
                color=DARK_BLUE,
                bold=bold,
            )
        elif token.startswith("["):
            label, url = token[1:].split("](", 1)
            add_hyperlink(paragraph, label, url[:-1])
        else:
            url = token[1:-1]
            add_hyperlink(paragraph, url, url)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:].replace(r"\*", "*"))
        set_run_font(run, bold=bold)


def add_shading(paragraph, fill: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = paragraph_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_callout(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    add_shading(paragraph, LIGHT_GRAY_FILL)
    add_inline_text(paragraph, text, bold=text.startswith(("重要", "注意")))


def next_numbering_id(numbering) -> tuple[int, int]:
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))
    ]
    return (max(abstract_ids, default=-1) + 1, max(num_ids, default=0) + 1)


def create_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id, num_id = next_numbering_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.extend((tabs, indent, spacing))
    level.extend((start, number_format, level_text, justification, paragraph_properties))
    abstract.append(level)
    numbering.append(abstract)
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), str(num_id))
    number_properties.extend((level, identifier))
    properties.append(number_properties)
    add_inline_text(paragraph, text)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def set_cell_width(cell, width_dxa: int) -> None:
    cell.width = Inches(width_dxa / 1440)
    cell_properties = cell._tc.get_or_add_tcPr()
    width = cell_properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        cell_properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def configure_table_geometry(table, column_widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(column_widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = table_properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_properties.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")
    table_layout = table_properties.first_child_found_in("w:tblLayout")
    if table_layout is None:
        table_layout = OxmlElement("w:tblLayout")
        table_properties.append(table_layout)
    table_layout.set(qn("w:type"), "fixed")

    margins = table_properties.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        table_properties.append(margins)
    for side, value in CELL_MARGIN_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in tuple(grid):
        grid.remove(child)
    for width_dxa in column_widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, column_widths):
            set_cell_width(cell, width_dxa)


def table_widths(column_count: int) -> list[int]:
    if column_count == 2:
        return [2700, CONTENT_WIDTH_DXA - 2700]
    base_width = CONTENT_WIDTH_DXA // column_count
    widths = [base_width] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows or not rows[0]:
        return
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError("Markdown 表格列数不一致")
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = table_widths(column_count)
    configure_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        cannot_split = OxmlElement("w:cantSplit")
        row_properties.append(cannot_split)
        if row_index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            row_properties.append(repeat_header)
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            add_inline_text(paragraph, value, bold=row_index == 0)
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), LIGHT_BLUE_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.08
    paragraph.paragraph_format.keep_together = False
    add_shading(paragraph, "F2F4F7")
    run = paragraph.add_run(code.rstrip())
    set_run_font(
        run,
        name=MONO_FONT,
        east_asia=EAST_ASIA_FONT,
        size=9,
        color=INK_BLUE,
    )


def add_image(
    doc: Document,
    source_path: Path,
    target: str,
    alt_text: str,
    *,
    max_width_inches: float = 6.15,
) -> None:
    if target.startswith("data:image/png;base64,"):
        try:
            image_bytes = base64.b64decode(target.split(",", 1)[1], validate=True)
        except (ValueError, base64.binascii.Error) as error:
            raise ValueError("Word 图片包含无效 Base64") from error
        if not image_bytes.startswith(export_guides_v2.PNG_SIGNATURE):
            raise ValueError("Word 图片数据不是 PNG")
        image_source = io.BytesIO(image_bytes)
        image = Image.from_file(image_source)
        picture_source = io.BytesIO(image_bytes)
        image_name = "embedded-v2-image"
    else:
        image_path = markdown_exporter.image_path_from_target(source_path, target)
        image = Image.from_file(str(image_path))
        picture_source = str(image_path)
        image_name = image_path.stem
    max_width = Inches(max_width_inches)
    max_height = Inches(7.15)
    width = image.width
    height = image.height
    scale = min(1.0, max_width / width, max_height / height)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(
        picture_source,
        width=int(width * scale),
        height=int(height * scale),
    )
    description = alt_text or image_name
    shape._inline.docPr.set("descr", description)
    shape._inline.docPr.set("title", description)


def is_special_line(lines: list[str], index: int) -> bool:
    stripped = lines[index].strip()
    if not stripped:
        return True
    if stripped.startswith(("```", "#", ">")):
        return True
    if IMAGE_PATTERN.match(stripped) or ORDERED_PATTERN.match(stripped) or BULLET_PATTERN.match(stripped):
        return True
    return (
        stripped.startswith("|")
        and index + 1 < len(lines)
        and TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip()) is not None
    )


def render_markdown_to_document(
    markdown: str,
    source_path: Path,
    *,
    is_parent: bool,
    edition: str = "v1",
    metadata: dict | None = None,
) -> Document:
    lines = markdown.splitlines()
    title_match = next(
        (HEADING_PATTERN.match(line.strip()) for line in lines if line.strip()),
        None,
    )
    if title_match is None or title_match.group(1) != "#":
        raise ValueError(f"教程缺少一级标题: {source_path.name}")
    title = title_match.group(2)
    doc = Document()
    if edition == "v2":
        word_v2.configure_document(doc, title, sys.modules[__name__])
        word_v2.add_editorial_opening(
            doc,
            title,
            metadata or {},
            sys.modules[__name__],
        )
    else:
        configure_document(doc, title, is_parent)
        add_cover(doc, title, is_parent)

    index = 0
    skipped_title = False
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            fence = stripped[:3]
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith(fence):
                code_lines.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError(f"代码块未闭合: {source_path.name}")
            add_code_block(doc, "\n".join(code_lines))
            index += 1
            continue

        heading = HEADING_PATTERN.match(stripped)
        if heading:
            markdown_level = len(heading.group(1))
            if markdown_level == 1 and not skipped_title:
                skipped_title = True
            else:
                word_level = min(max(markdown_level - 1, 1), 3)
                paragraph = doc.add_paragraph(style=f"Heading {word_level}")
                add_inline_text(paragraph, heading.group(2))
            index += 1
            continue

        image_match = IMAGE_PATTERN.match(stripped)
        if image_match:
            add_image(
                doc,
                source_path,
                image_match.group(2),
                image_match.group(1),
                max_width_inches=5.0 if edition == "v2" else 6.15,
            )
            index += 1
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and TABLE_SEPARATOR_PATTERN.match(lines[index + 1].strip())
        ):
            rows = [parse_table_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(parse_table_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue

        ordered_match = ORDERED_PATTERN.match(stripped)
        bullet_match = BULLET_PATTERN.match(stripped)
        if ordered_match or bullet_match:
            ordered = ordered_match is not None
            num_id = create_numbering(doc, ordered=ordered)
            matcher = ORDERED_PATTERN if ordered else BULLET_PATTERN
            while index < len(lines):
                item_match = matcher.match(lines[index].strip())
                if item_match is None:
                    break
                item_text = item_match.group(1)
                index += 1
                continuation: list[str] = []
                while (
                    index < len(lines)
                    and lines[index].strip()
                    and not is_special_line(lines, index)
                    and lines[index][:1].isspace()
                ):
                    continuation.append(lines[index].strip())
                    index += 1
                if continuation:
                    item_text = f"{item_text} {' '.join(continuation)}"
                add_list_item(doc, item_text, num_id)
                while index < len(lines) and not lines[index].strip():
                    index += 1
            continue

        if stripped.startswith(">"):
            add_callout(doc, stripped.lstrip("> "))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not is_special_line(lines, index):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph_text = " ".join(paragraph_lines)
        if paragraph_text.startswith(("重要：", "注意：", "提示：")):
            add_callout(doc, paragraph_text)
        else:
            style = "Caption" if paragraph_text.startswith(("图：", "图 ")) else None
            paragraph = doc.add_paragraph(style=style)
            if style is None:
                paragraph.paragraph_format.widow_control = True
            add_inline_text(paragraph, paragraph_text)
    return doc


def _normalize_rfonts_attributes(
    payload: bytes,
    *,
    body_font: str | None = None,
    east_asia_font: str | None = None,
) -> bytes:
    body_font_bytes = body_font.encode() if body_font else None
    east_asia_font_bytes = east_asia_font.encode() if east_asia_font else None

    def replace_rfonts(element_match: re.Match[bytes]) -> bytes:
        def replace_attribute(attribute_match: re.Match[bytes]) -> bytes:
            attribute = attribute_match.group(1)
            quote = attribute_match.group(2)
            value = attribute_match.group(3)
            replacement = value
            if (
                body_font_bytes
                and attribute in {b"ascii", b"hAnsi"}
                and value == BODY_FONT.encode()
            ):
                replacement = body_font_bytes
            if (
                east_asia_font_bytes
                and attribute in {b"eastAsia", b"cs"}
                and value == EAST_ASIA_FONT.encode()
            ):
                replacement = east_asia_font_bytes
            return b"w:" + attribute + b"=" + quote + replacement + quote

        return RFONTS_ATTRIBUTE_PATTERN.sub(
            replace_attribute,
            element_match.group(0),
        )

    return RFONTS_ELEMENT_PATTERN.sub(replace_rfonts, payload)


def normalize_docx_package(
    raw_docx: bytes,
    *,
    body_font: str | None = None,
    east_asia_font: str | None = None,
) -> bytes:
    source_buffer = io.BytesIO(raw_docx)
    destination_buffer = io.BytesIO()
    with zipfile.ZipFile(source_buffer) as source_archive, zipfile.ZipFile(
        destination_buffer,
        mode="w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as destination_archive:
        for source_info in sorted(source_archive.infolist(), key=lambda item: item.filename):
            normalized_info = zipfile.ZipInfo(source_info.filename, FIXED_PACKAGE_TIME)
            normalized_info.compress_type = zipfile.ZIP_DEFLATED
            normalized_info.create_system = 3
            normalized_info.external_attr = source_info.external_attr
            payload = source_archive.read(source_info.filename)
            if source_info.filename.endswith(".xml") and (
                body_font or east_asia_font
            ):
                payload = _normalize_rfonts_attributes(
                    payload,
                    body_font=body_font,
                    east_asia_font=east_asia_font,
                )
            destination_archive.writestr(normalized_info, payload)
    return destination_buffer.getvalue()


def render_document(guide: markdown_exporter.GuideSpec) -> bytes:
    source_path = markdown_exporter.SOURCE_DIR / guide.source_name
    markdown = markdown_exporter.render_guide_markdown(guide)
    document = render_markdown_to_document(
        markdown,
        source_path,
        is_parent=guide is markdown_exporter.GUIDES[0],
    )
    output_buffer = io.BytesIO()
    document.save(output_buffer)
    return normalize_docx_package(output_buffer.getvalue())

def rendered_documents() -> tuple[tuple[str, bytes], ...]:
    return tuple(
        (str(Path(guide.output_name).with_suffix(".docx")), render_document(guide))
        for guide in markdown_exporter.GUIDES
    )


def rendered_v2_documents() -> tuple[tuple[str, bytes], ...]:
    return word_v2.rendered_documents(sys.modules[__name__])


def main() -> int:
    args = parse_args()
    default_output = (
        export_guides_v2.DEFAULT_WORD_OUTPUT_DIR
        if args.edition == "v2"
        else DEFAULT_OUTPUT_DIR
    )
    output_dir = args.output_dir or default_output
    try:
        if args.edition == "v2":
            exports = rendered_v2_documents()
            if args.check:
                errors = export_guides_v2.check_export_tree(output_dir, exports)
                if errors:
                    print("\n".join(errors), file=sys.stderr)
                    return 1
                print(f"校验通过：{len(exports)} 份 V2 Word 教程与源稿一致。")
                return 0
            export_guides_v2.atomic_export_tree(output_dir, exports)
            print(f"已生成 {len(exports)} 份图片内嵌 V2 Word 教程：{output_dir}")
            return 0

        output_dir = output_dir.resolve()
        exports = rendered_documents()
        if args.check:
            errors = check_documents(output_dir, exports)
            if errors:
                print("\n".join(errors), file=sys.stderr)
                return 1
            print(f"校验通过：{len(exports)} 份 Word 教程与源稿一致。")
            return 0
        export_documents(output_dir, exports)
    except (OSError, UnicodeError, ValueError, zipfile.BadZipFile) as error:
        print(f"生成失败：{error}", file=sys.stderr)
        return 1
    print(f"已生成 {len(exports)} 份图片内嵌 Word 教程：{output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

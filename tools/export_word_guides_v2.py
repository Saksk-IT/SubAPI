"""V2 Word 的紧凑开篇、中文字体回退与九份文档渲染。"""

from __future__ import annotations

import io
import re
from types import ModuleType

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

if __package__:
    from . import export_guides_v2
else:
    import export_guides_v2


EAST_ASIA_FONT = "Hiragino Sans GB"
CJK_PATTERN = re.compile(r"[\u3400-\u9fff]")


def configure_document(doc, title: str, renderer: ModuleType) -> None:
    renderer.configure_document(doc, title, is_parent=False)
    header = doc.sections[0].header.paragraphs[0]
    header.clear()
    header_run = header.add_run(f"{export_guides_v2.BRAND} · {title}")
    renderer.set_run_font(header_run, size=8.5, color=renderer.MUTED_GRAY)
    properties = doc.core_properties
    properties.subject = "AI 客户端配置教程"
    properties.author = export_guides_v2.BRAND
    properties.last_modified_by = export_guides_v2.BRAND
    doc.styles["Caption"].paragraph_format.space_after = Pt(6)


def add_editorial_opening(
    doc,
    title: str,
    metadata: dict,
    renderer: ModuleType,
) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    kicker_run = kicker.add_run(export_guides_v2.BRAND)
    renderer.set_run_font(
        kicker_run,
        size=9.5,
        color=renderer.ACCENT_BLUE,
        bold=True,
    )

    title_paragraph = doc.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(5)
    title_paragraph.paragraph_format.keep_with_next = True
    title_run = title_paragraph.add_run(title)
    renderer.set_run_font(
        title_run,
        size=21,
        color=renderer.INK_BLUE,
        bold=True,
    )

    summary = str(metadata.get("summary", "")).strip()
    if summary:
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.paragraph_format.space_after = Pt(5)
        summary_paragraph.paragraph_format.keep_with_next = True
        summary_run = summary_paragraph.add_run(summary)
        renderer.set_run_font(
            summary_run,
            size=11.5,
            color=renderer.MUTED_GRAY,
        )

    platforms = " / ".join(str(value) for value in metadata.get("platforms", ()))
    meta_text = " · ".join(
        value
        for value in (
            f"预计 {metadata.get('duration', '')}" if metadata.get("duration") else "",
            f"难度 {metadata.get('difficulty', '')}" if metadata.get("difficulty") else "",
            platforms,
            f"更新 {metadata.get('updatedAt', '')}" if metadata.get("updatedAt") else "",
        )
        if value
    )
    meta_paragraph = doc.add_paragraph()
    meta_paragraph.paragraph_format.space_after = Pt(10)
    meta_paragraph.paragraph_format.keep_with_next = True
    meta_run = meta_paragraph.add_run(meta_text)
    renderer.set_run_font(
        meta_run,
        size=9,
        color=renderer.ACCENT_BLUE,
        bold=True,
    )


def apply_cjk_fallback(doc) -> None:
    roots = [doc.element]
    for section in doc.sections:
        roots.extend((section.header._element, section.footer._element))
    for root in roots:
        for run in root.iter(qn("w:r")):
            text = "".join(node.text or "" for node in run.iter(qn("w:t")))
            if not CJK_PATTERN.search(text):
                continue
            properties = run.get_or_add_rPr()
            fonts = properties.find(qn("w:rFonts"))
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                properties.insert(0, fonts)
            for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{attribute}"), EAST_ASIA_FONT)


def render_document(
    renderer: ModuleType,
    guide: export_guides_v2.GuideSpec,
    markdown: str,
    metadata: dict,
) -> bytes:
    document = renderer.render_markdown_to_document(
        markdown,
        export_guides_v2.CONTENT_DIR / guide.source_name,
        is_parent=guide is export_guides_v2.GUIDES[0],
        edition="v2",
        metadata=metadata,
    )
    apply_cjk_fallback(document)
    output_buffer = io.BytesIO()
    document.save(output_buffer)
    return renderer.normalize_docx_package(
        output_buffer.getvalue(),
        body_font="Calibri",
        east_asia_font=EAST_ASIA_FONT,
    )


def rendered_documents(renderer: ModuleType) -> tuple[tuple[str, bytes], ...]:
    bodies = export_guides_v2.load_verified_bodies()
    media = export_guides_v2.load_media()
    metadata = export_guides_v2.load_manifest_metadata()
    return tuple(
        (
            guide.output_name(".docx"),
            render_document(
                renderer,
                guide,
                export_guides_v2.render_markdown(bodies[guide.source_name], media),
                metadata[guide.source_name],
            ),
        )
        for guide in export_guides_v2.GUIDES
    )
